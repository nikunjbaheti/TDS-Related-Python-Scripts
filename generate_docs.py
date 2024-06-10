import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx2pdf import convert
import pandas as pd

# Load the data
data = pd.read_csv('data.csv')

# Get the current directory where the script is located
current_directory = os.path.dirname(os.path.abspath(__file__))

# Path to the template
template_path = os.path.join(current_directory, 'template.docx')

# Function to replace placeholders and set font in the document
def replace_placeholder(doc, placeholder, replacement):
    for p in doc.paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                run.text = run.text.replace(placeholder, replacement)
                run.font.name = 'Aptos Narrow'
                run.font.size = Pt(18)
                run.bold = True
            p.style.font.name = 'Aptos Narrow'
            p.style.font.size = Pt(18)
            p.style.font.bold = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, replacement)

# Generate documents
for index, row in data.iterrows():
    # Create a copy of the template
    doc = Document(template_path)
    
    # Replace placeholders with actual data
    replace_placeholder(doc, '{{Employee Name}}', row['Employee Name'])
    replace_placeholder(doc, '{{Employee PAN}}', row['Employee PAN'])
    replace_placeholder(doc, '{{Assessment Year}}', row['Assessment Year'])
    replace_placeholder(doc, '{{Financial Year}}', row['Financial Year'])
    
    # Define the output path for the .docx file
    output_docx_path = os.path.join(current_directory, f'{row["Employee PAN"]}.docx')
    
    # Save the new document
    doc.save(output_docx_path)
    
    # Convert the .docx file to .pdf
    output_pdf_path = os.path.join(current_directory, f'{row["Employee PAN"]}.pdf')
    convert(output_docx_path, output_pdf_path)
    
    # Delete the .docx file
    os.remove(output_docx_path)

print("Documents generated and converted to PDF successfully!")
